#!/usr/bin/env python
# coding: utf-8

# In[2]:


"""Code to convert Wrod document .docx to txt file."""

from docx import Document

def wordToTxt(filename):
    document = Document(filename)
    f=open("/Users/pradeep/Documents/Project/resumes/Train/BigData/check2.txt", "w+")
    c = 0
    data = []
    keys = None
    for para in document.paragraphs:
        f.write(para.text)
        f.write("\n")

    tableData = document.tables
    
    for table in tableData:
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            h = str(tuple(text))
            f.write(h)

        f.write("\n")

    f.close()

def main(file):
    wordToTxt(file)
    
if __name__ == '__main__':
    main('151039001_Shreehari R.docx')


# In[3]:


import os
from docx import Document
fileNo = 1
def wordToTxt(file, fileNo):
    filename = "/Users/pradeep/Documents/Project/resumes/Embedded Systems/" + file
    document = Document(filename)
    print(document)
    
    resumeStorage = "/Users/pradeep/Documents/Project/resumes/Train/ES/" + "txt" + str(fileNo)
    f=open(resumeStorage, "w+")
    c = 0
    data = []
    keys = None
    for para in document.paragraphs:
        f.write(para.text)
        f.write("\n")

    tableData = document.tables
    
    for table in tableData:
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            h = str(tuple(text))
            f.write(h)

        f.write("\n")
    f.close()
#     fileNo += 1
    return True
        
count = 0
path = "/Users/pradeep/Documents/Project/CVs/2018/bigData/Big Data"
for root, dirs, files in os.walk(path):  
    for filename in files:
        if ".docx" in filename:
            count += 1
            wordToTxt(filename, count)
#             print(filename)
#         count += 1
#         print(filename[-4:])
#         break;

print(count)


# In[11]:


count = 0
path = "/Users/pradeep/Documents/Project/resumes/Embedded Systems"
for root, dirs, files in os.walk(path):  
    for filename in files:
        if ".doc" in filename:
            count += 1
#             wordToTxt(filename)
            print(count, filename)
#         count += 1
#         print(filename[-4:])
#         break;

print(count)


# In[2]:


import docx2txt
text = docx2txt.process("161046013-Ravali.docx")
k = open("j.txt", "w+")
k.write(text)
k.close()


# In[11]:


import os
import docx2txt
# from docx import Document
fileNo = 1
def wordToTxt(file, fileNo):
    filename = "/Users/pradeep/Documents/Project/CVs/2018/VLSI-Resumes/" + file
    resumeStorage = "/Users/pradeep/Documents/Project/resumes/Test Data/resume.vlsi/" + "txt" + str(fileNo)
    
    text = docx2txt.process(filename)
    k = open(resumeStorage, "w+")
#     k = open("j.txt", "w+")
    k.write(text)
    k.close()
    
    return True
        
count = 0
path = "/Users/pradeep/Documents/Project/CVs/2018/VLSI-Resumes/"
for root, dirs, files in os.walk(path):  
    for filename in files:
        if ".docx" in filename:
            count += 1
            wordToTxt(filename, count)
#             print(filename)
#         count += 1
#         print(filename[-4:])
#         break;

print(count)


# In[46]:


import re
def extract(text):
    pattern = re.compile(r'skills[\s\S]*?personal')
    x = re.findall(pattern, text)
    print(x)


# In[47]:


df = open("/Users/pradeep/Documents/Project/resumes/Training Data/BigData/txt34", "rb")
text = df.read()
text = str(text)
text = str(text.lower())
extract(text)
# print(text)
df.close()


# In[ ]:




